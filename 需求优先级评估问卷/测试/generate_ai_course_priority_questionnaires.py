from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "需求优先级评估问卷"
SRS_REF = "AI课程教辅系统_SRS_V1.3.docx"
PROJECT_NAME = "AI课程教辅系统"
PROJECT_CODE = "SRS-AICS-2026-01"
DATE_TEXT = "2026年5月31日"


MODULES = [
    {
        "id": "M01",
        "name": "公共入口、登录注册与角色权限",
        "kano": "E",
        "desc": "提供系统入口、品牌与课程定位、账号登录、教师/学生注册、个人中心、退出登录和基于角色的访问控制。",
        "ref": "对应SRS 3.2.1 UC-00~UC-03、UC-P01~UC-P03。",
    },
    {
        "id": "M02",
        "name": "教学资源上传、管理与AI解析",
        "kano": "E",
        "desc": "教师上传PPT/PDF/Word/TXT等课程资源，系统完成文本提取、切片、向量化、知识点抽取、题目生成和解析状态跟踪。",
        "ref": "对应SRS 3.2.2 UC-T01、UC-T02。",
    },
    {
        "id": "M03",
        "name": "知识图谱审核、发布与学习导航",
        "kano": "E",
        "desc": "教师审核AI抽取的知识点、前置关系和掌握要求，发布后学生通过知识图谱查看课程结构、薄弱节点和推荐资料。",
        "ref": "对应SRS 3.2.2 UC-T03、3.2.3 UC-S01。",
    },
    {
        "id": "M04",
        "name": "智能备课与教学内容生成",
        "kano": "N",
        "desc": "教师通过AI助手生成备课方案、课堂提问、教案、PPT或教学视频脚本，支持按章节和教学目标持续追问。",
        "ref": "对应SRS 3.2.2 UC-T04。",
    },
    {
        "id": "M05",
        "name": "自动组卷、试题审核与发布",
        "kano": "E",
        "desc": "教师配置题型、数量、分值、难度和薄弱知识点覆盖策略，AI生成试卷草稿，教师审核来源、答案、解析和置信度后发布。",
        "ref": "对应SRS 3.2.2 UC-T05、UC-T06。",
    },
    {
        "id": "M06",
        "name": "学生在线练习与提交反馈",
        "kano": "E",
        "desc": "学生完成教师发布或AI推荐的专项练习，系统记录答题、标记、倒计时、得分、正确率、错题和薄弱知识点。",
        "ref": "对应SRS 3.2.3 UC-S03。",
    },
    {
        "id": "M07",
        "name": "错题本与错题变式训练",
        "kano": "N",
        "desc": "系统按章节、知识点、题型和错误次数管理错题，基于原错题生成相似情境的变式题，帮助学生巩固薄弱知识点。",
        "ref": "对应SRS 3.2.3 UC-S04、UC-S05。",
    },
    {
        "id": "M08",
        "name": "能力雷达图与个性化学习路径",
        "kano": "N",
        "desc": "系统根据练习成绩、错题、章节掌握度和班级平均水平生成能力雷达图、薄弱项分析和分步学习路径。",
        "ref": "对应SRS 3.2.3 UC-S02、UC-S06。",
    },
    {
        "id": "M09",
        "name": "AI课程问答与RAG来源引用",
        "kano": "E",
        "desc": "学生基于课程知识库提问，系统通过RAG检索生成回答，并显示引用来源、相关知识点、推荐练习和学习免责声明。",
        "ref": "对应SRS 3.2.3 UC-S07、3.4.2 RAG与来源约束。",
    },
    {
        "id": "M10",
        "name": "课堂数据分析与学生预警",
        "kano": "N",
        "desc": "教师查看班级完成率、平均分、知识点热力图、章节趋势、高频错题和风险学生，并推送练习或发送提醒。",
        "ref": "对应SRS 3.2.2 UC-T07、UC-T08。",
    },
    {
        "id": "M11",
        "name": "后台用户、课程与知识库管理",
        "kano": "E",
        "desc": "管理员维护教师、学生、管理员账号，管理课程、班级、授课关系、知识库文档、索引状态、切片数量和重建索引任务。",
        "ref": "对应SRS 3.2.4 UC-A01~UC-A03。",
    },
    {
        "id": "M12",
        "name": "模型服务监控、权限安全与系统日志",
        "kano": "E",
        "desc": "管理员监控RAG平均响应时间、AI请求成功率、并发量、错误日志，配置权限边界，导出系统日志并审计危险操作。",
        "ref": "对应SRS 3.2.4 UC-A04~UC-A06、3.6.1安全性。",
    },
]


ADMIN_SCORES = {
    "M01": (8, 8, 5, 4),
    "M02": (7, 6, 6, 5),
    "M03": (8, 8, 6, 5),
    "M04": (6, 4, 6, 5),
    "M05": (7, 7, 5, 4),
    "M06": (6, 5, 5, 4),
    "M07": (5, 4, 5, 4),
    "M08": (6, 5, 4, 4),
    "M09": (8, 7, 6, 5),
    "M10": (7, 6, 5, 4),
    "M11": (9, 8, 5, 4),
    "M12": (9, 9, 5, 5),
}

TEACHER_SCORES = {
    "M01": ("8", "8"),
    "M02": ("9", "9"),
    "M03": ("9", "9"),
    "M04": ("8", "6"),
    "M05": ("9", "8"),
    "M06": ("8", "8"),
    "M07": ("6", "4"),
    "M08": ("9", "7"),
    "M09": ("6", "5"),
    "M10": ("8", "7"),
    "M11": ("不适用", "不适用"),
    "M12": ("5", "4"),
}

STUDENT_SCORES = {
    "M01": ("8", "7"),
    "M02": ("5", "4"),
    "M03": ("8", "7"),
    "M04": ("不适用", "不适用"),
    "M05": ("不适用", "不适用"),
    "M06": ("9", "9"),
    "M07": ("9", "8"),
    "M08": ("9", "8"),
    "M09": ("9", "9"),
    "M10": ("6", "5"),
    "M11": ("不适用", "不适用"),
    "M12": ("6", "5"),
}


def set_run_font(run, size: int | None = None, bold: bool | None = None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def add_p(doc: Document, text: str = "", bold: bool = False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold)
    return p


def heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def set_cell(cell, text: str, bold: bool = False):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            set_run_font(run, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], str(value))
    doc.add_paragraph()
    return table


def add_cover(doc: Document, person: dict):
    for text, size, bold in [
        (f"{PROJECT_NAME}（{PROJECT_CODE}）", 16, True),
        ("功能需求优先级评估问卷", 18, True),
        ("—— 基于质量功能调配（QFD）方法 ——", 12, False),
    ]:
        p = add_p(doc, text, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in p.runs:
            run.font.size = Pt(size)

    add_p(doc)
    for line in [
        f"评估人姓名：{person['name']}",
        f"评估人角色：{person['role']}",
        f"所属用户群：{person['group']}",
        f"电子邮箱：{person['email']}",
        f"填表日期：{DATE_TEXT}",
        f"编写依据：{SRS_REF}、开发原型1.1、用例图与用例文档",
    ]:
        add_p(doc, line)


def add_method(doc: Document, perspective: str, cost_note: str):
    heading(doc, "一、评估方法说明", 1)
    add_p(
        doc,
        f"本问卷采用质量功能调配（Quality Function Deployment, QFD）方法，对{PROJECT_NAME}在{perspective}下的功能需求进行优先级量化评估。QFD的核心思想是将产品功能与用户价值、缺失损失和实现代价联系起来，为版本范围和迭代顺序提供依据。",
    )
    heading(doc, "1.1 需求的三种类型（Kano模型视角）", 2)
    add_p(doc, "（1）期望需求（E）：用户认为系统理所当然应具备的功能，缺失会直接导致不满意。")
    add_p(doc, "（2）普通需求（N）：功能越完善满意度越高，是系统竞争力和使用效率的重要来源。")
    add_p(doc, "（3）兴奋需求（X）：用户未必明确要求，但提供后能显著提升体验或形成亮点。")
    heading(doc, "1.2 QFD优先级评估维度（1~9分制）", 2)
    add_p(doc, "（1）相对利益：该功能实现后将给本角色带来多大价值？1=几乎无价值，5=中等价值，9=极大价值。")
    add_p(doc, "（2）相对损失：该功能如不包含在首轮实现范围内将带来多大损失？1=基本无损失，5=中等损失，9=损失极为严重。")
    add_p(doc, "（3）实现成本：实现该功能预计需要多少开发资源？1=成本极低，5=中等成本，9=成本极高。")
    add_p(doc, "（4）技术风险：实现该功能的技术难度和不确定性有多高？1=基本无风险，5=中等风险，9=风险极高。")
    heading(doc, "1.3 优先级计算公式", 2)
    add_p(doc, "每个功能模块的综合优先级得分按以下公式计算：")
    add_p(doc, "                    相对利益 + 相对损失")
    add_p(doc, "优先级得分  =  ───────────────────")
    add_p(doc, "                    实现成本 + 技术风险")
    add_p(doc, f"优先级得分越高，说明在单位实现代价下能获得越大的用户价值。{cost_note}")


def add_background(doc: Document, background: str, pain_points: list[str]):
    heading(doc, "二、个人使用背景", 1)
    add_p(doc, background)
    heading(doc, "三、现有痛点", 1)
    for idx, item in enumerate(pain_points, start=1):
        add_p(doc, f"{idx}. {item}")


def admin_rows():
    rows = []
    for module in MODULES:
        b, p, c, r = ADMIN_SCORES[module["id"]]
        score = (b + p) / (c + r)
        rows.append(
            [
                module["id"],
                module["name"],
                module["kano"],
                module["desc"],
                module["ref"],
                b,
                p,
                c,
                r,
                f"{score:.2f}",
            ]
        )
    return rows


def role_rows(scores: dict[str, tuple[str, str]], perspective: str):
    rows = []
    for module in MODULES:
        benefit, penalty = scores[module["id"]]
        rows.append(
            [
                module["id"],
                module["name"],
                module["kano"],
                module["desc"],
                f"{module['ref']}；请从{perspective}判断该模块是否必须进入首轮实现。",
                benefit,
                penalty,
            ]
        )
    return rows


def add_qfd_table(doc: Document, person: dict):
    heading(doc, "四、QFD优先级评估表", 1)
    add_p(doc, person["table_intro"])
    if person["kind"] == "admin":
        add_table(
            doc,
            [
                "编号",
                "功能模块",
                "Kano\n类型",
                "功能说明",
                "SRS对应说明",
                "相对利益\n(1~9)",
                "相对损失\n(1~9)",
                "实现成本\n(1~9)",
                "技术风险\n(1~9)",
                "优先级得分",
            ],
            admin_rows(),
        )
    else:
        add_table(
            doc,
            [
                "编号",
                "功能模块",
                "Kano\n类型",
                f"功能说明（{person['short_role']}视角）",
                "评分参考说明",
                "相对利益\n(1~9)",
                "相对损失\n(1~9)",
            ],
            role_rows(person["scores"], person["short_role"]),
        )
        add_p(doc, "注：实现成本、技术风险和综合优先级得分由项目组在回收问卷后统一计算。")


def add_suggestions(doc: Document, person: dict):
    heading(doc, "五、整体评价与建议", 1)
    for para in person["suggestions"]:
        add_p(doc, para)


def add_signature(doc: Document, person: dict):
    heading(doc, "六、评估人签署", 1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["评估人签名", person["signature"]],
            ["日期", DATE_TEXT],
            ["备注", f"本问卷评分将作为{PROJECT_NAME}需求优先级分析的输入，直接用于SRS V1.3需求基线和首轮实现范围决策。"],
        ],
    )


def build_doc(person: dict, path: Path):
    doc = Document()
    style_doc(doc)
    add_cover(doc, person)
    add_method(doc, person["perspective"], person["cost_note"])
    add_background(doc, person["background"], person["pain_points"])
    add_qfd_table(doc, person)
    add_suggestions(doc, person)
    add_signature(doc, person)
    doc.save(path)


PEOPLE = [
    {
        "kind": "admin",
        "name": "吴易骏",
        "role": "系统管理员代表（平台治理与运维方）",
        "short_role": "管理员",
        "group": "系统管理方",
        "email": "32301304@stu.hzcu.edu.cn",
        "signature": "吴易骏",
        "perspective": "系统管理员与平台运维",
        "cost_note": "管理员代表可结合运维、权限、日志、模型监控和系统治理经验，对实现成本与技术风险给出初步估计。",
        "background": "作为系统管理员代表，我主要从平台可维护性、权限安全、课程数据隔离、RAG知识库治理、模型服务稳定性和系统审计的角度评估需求优先级。AI课程教辅系统不仅要支持教学闭环，还要保证账号、课程、资源、知识库、模型调用和日志等后台能力可管、可查、可追溯。",
        "pain_points": [
            "课程资源、知识库文档、AI解析任务和题目草稿之间存在复杂关联，如果没有后台治理能力，后续维护成本会很高。",
            "AI问答和自动组卷依赖RAG检索，必须能监控向量化状态、切片数、响应时间、成功率和错误日志。",
            "教师、学生、管理员的权限边界必须清晰，跨课程访问、越权修改审核结论、无来源AI回答都需要被系统阻止。",
            "删除文档、重建索引、禁用账号、发布试卷等操作影响范围大，必须有二次确认和审计日志。",
        ],
        "table_intro": "以下从管理员代表的角度，对12个功能模块逐一在四个维度上评分。各模块已标注Kano需求分类（E=期望需求，N=普通需求，X=兴奋需求）。",
        "suggestions": [
            "从管理方视角，首轮实现应优先保证M11后台用户、课程与知识库管理，M12模型服务监控、权限安全与系统日志，以及M09 AI课程问答与RAG来源引用。这些模块决定系统是否可持续运行和可审计。",
            "M02资源上传与AI解析、M03知识图谱审核、M05试题审核与发布虽然主要由教师使用，但对后台索引、内容状态和日志追踪有直接影响，也应纳入首轮稳定范围。",
            "建议将高风险操作全部纳入审计范围，并在SRS验收中增加权限隔离、来源引用、索引重建、日志导出和异常告警的检查项。",
        ],
    },
    {
        "kind": "teacher",
        "name": "杨枨",
        "role": "教师代表",
        "short_role": "教师",
        "group": "授课教师",
        "email": "yangc@stu.hzcu.edu.cn",
        "signature": "杨枨",
        "perspective": "授课教师",
        "cost_note": "教师代表主要评估相对利益和相对损失；实现成本和技术风险由项目组结合技术方案统一评估。",
        "background": "作为《软件需求分析原理与实践》的授课教师代表，我关注系统能否减少备课、资源整理、知识点梳理、出题、批改分析和学生干预的重复劳动。教师端应帮助我把课程资料转化为可审核的知识图谱、题目、练习和学习分析结果，同时保证AI生成内容可核验、可修改、可发布。",
        "pain_points": [
            "课程资料分散在PPT、Word、PDF和案例文件中，人工整理知识点和前置关系耗时较长。",
            "备课、出题和试卷审核常常重复劳动，尤其是要兼顾知识点覆盖率、题型结构、难度分布和薄弱知识点。",
            "学生做完练习后，教师难以及时看到班级层面的知识点掌握热力图和个人风险学生。",
            "AI生成内容如果没有来源引用、置信度和人工审核流程，教师无法放心直接用于教学。",
        ],
        "table_intro": "请从教师用户的角度，对以下12个功能模块逐一在“相对利益”“相对损失”两个维度上打分（1~9分）。如该模块不在您的关注范围内，请在对应行标注“不适用”。",
        "scores": TEACHER_SCORES,
        "suggestions": [
            "从教师视角，M02教学资源上传、管理与AI解析，M03知识图谱审核与发布，M05自动组卷、试题审核与发布是首轮必须实现的核心功能。",
            "M08能力雷达图与个性化学习路径、M10课堂数据分析与学生预警可以显著提升教学反馈效率，也建议进入首轮可演示范围。",
            "M04智能备课可以作为普通需求逐步增强，首轮至少应支持章节备课方案、课堂提问和教案草稿生成。",
        ],
    },
    {
        "kind": "student",
        "name": "蒋长麟",
        "role": "学生代表（在课学生）",
        "short_role": "学生",
        "group": "在课学生",
        "email": "32301294@stu.hzcu.edu.cn",
        "signature": "蒋长麟",
        "perspective": "在课学生",
        "cost_note": "学生代表主要评估相对利益和相对损失；实现成本和技术风险由项目组结合技术方案统一评估。",
        "background": "作为在课学生代表，我主要关注系统能否帮助我理解《软件需求分析原理与实践》的知识结构、找到薄弱知识点、完成练习、复盘错题，并通过可信的AI课程问答获得带来源的解释。学生端应减少盲目刷题和无效搜索，让学习路径更清楚。",
        "pain_points": [
            "课程章节、知识点和题目之间的关系不够直观，复习时不知道该先补哪个知识点。",
            "练习结束后只看到分数帮助有限，更需要知道错题对应的知识点、薄弱章节和下一步学习建议。",
            "错题如果只是静态保存，容易出现记答案而非真正理解的问题，希望能生成变式题继续训练。",
            "向AI提问时最担心回答不准确，因此需要看到课程资料引用来源和推荐练习。",
        ],
        "table_intro": "请从学生用户的角度，对以下12个功能模块逐一在“相对利益”“相对损失”两个维度上打分（1~9分）。如该模块不在您的关注范围内，请在对应行标注“不适用”。",
        "scores": STUDENT_SCORES,
        "suggestions": [
            "从学生视角，M06在线练习与提交反馈、M07错题本与错题变式训练、M08能力雷达图与学习路径、M09 AI课程问答是最直接影响学习体验的核心模块。",
            "M03知识图谱导航也应进入首轮范围，因为它决定学生能否理解课程知识结构并定位薄弱节点。",
            "后台管理类模块对学生不是直接功能，但M12中的权限安全、日志和数据保护会影响个人学习数据安全，建议作为系统底线能力保留。",
        ],
    },
]


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = []
    for person in PEOPLE:
        filename = f"需求优先级评估问卷-{person['short_role']}代表-{person['name']}-AI课程教辅系统.docx"
        path = OUT_DIR / filename
        build_doc(person, path)
        outputs.append(path)
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
